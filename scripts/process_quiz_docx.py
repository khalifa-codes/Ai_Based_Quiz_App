#!/usr/bin/env python3
"""
Quiz DOCX Processor
Processes uploaded DOCX quiz files, extracts questions, and converts to JSON format.

Usage:
    python process_quiz_docx.py <docx_file_path> [output_dir]

Returns:
    Exit code 0: Success
    Exit code 1: File not found
    Exit code 2: Invalid format
    Exit code 3: Extraction error
    Exit code 4: JSON conversion error
    Exit code 5: Save error
"""

import sys
import os
import json
import logging
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx library not installed. Install with: pip install python-docx")
    sys.exit(1)


# Configure logging
def setup_logging(log_dir: str = "storage/logs") -> logging.Logger:
    """Setup logging configuration"""
    os.makedirs(log_dir, exist_ok=True)
    log_file = os.path.join(log_dir, f"quiz_processor_{datetime.now().strftime('%Y%m%d')}.log")
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file),
            logging.StreamHandler(sys.stderr)
        ]
    )
    return logging.getLogger(__name__)


logger = setup_logging()


class QuizProcessor:
    """Process DOCX quiz files and convert to JSON"""
    
    def __init__(self, docx_path: str, output_dir: str = "storage/quiz_json"):
        self.docx_path = docx_path
        self.output_dir = output_dir
        self.document = None
        self.questions = []
        self.metadata = {}
        
        # Ensure output directory exists
        os.makedirs(self.output_dir, exist_ok=True)
    
    def validate_file(self) -> bool:
        """Validate that the DOCX file exists and is readable"""
        try:
            if not os.path.exists(self.docx_path):
                logger.error(f"File not found: {self.docx_path}")
                return False
            
            if not self.docx_path.lower().endswith('.docx'):
                logger.error(f"Invalid file type. Expected .docx, got: {os.path.splitext(self.docx_path)[1]}")
                return False
            
            # Try to open the document
            self.document = Document(self.docx_path)
            logger.info(f"Successfully opened DOCX file: {self.docx_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error validating file: {str(e)}")
            return False
    
    def extract_metadata(self) -> Dict[str, Any]:
        """Extract quiz metadata from document"""
        metadata = {
            'title': 'Untitled Quiz',
            'subject': '',
            'description': '',
            'total_marks': 0,
            'duration': 3600,  # Default 1 hour in seconds
            'created_at': datetime.now().isoformat(),
            'source_file': os.path.basename(self.docx_path)
        }
        
        try:
            # Check first few paragraphs for metadata
            for i, paragraph in enumerate(self.document.paragraphs[:10]):
                text = paragraph.text.strip()
                
                # Extract title (usually first non-empty paragraph)
                if i == 0 and text and not metadata['title']:
                    metadata['title'] = text
                
                # Look for metadata patterns
                if ':' in text:
                    parts = text.split(':', 1)
                    key = parts[0].strip().lower()
                    value = parts[1].strip()
                    
                    if 'subject' in key or 'course' in key:
                        metadata['subject'] = value
                    elif 'duration' in key or 'time' in key:
                        # Parse duration (e.g., "60 minutes" or "1 hour")
                        duration_match = re.search(r'(\d+)', value)
                        if duration_match:
                            minutes = int(duration_match.group(1))
                            metadata['duration'] = minutes * 60
                    elif 'marks' in key or 'total' in key:
                        marks_match = re.search(r'(\d+)', value)
                        if marks_match:
                            metadata['total_marks'] = int(marks_match.group(1))
            
            logger.info(f"Extracted metadata: {metadata}")
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}. Using defaults.")
            return metadata
    
    def is_question_start(self, text: str) -> bool:
        """Check if a paragraph is the start of a question"""
        # Patterns: "Q1.", "Question 1:", "1.", etc.
        patterns = [
            r'^Q\d+[\.:]?\s*',
            r'^Question\s+\d+[\.:]?\s*',
            r'^\d+[\.\)]\s+',
            r'^\d+\.\s+'
        ]
        
        text_stripped = text.strip()
        for pattern in patterns:
            # Use case-insensitive matching
            if re.match(pattern, text_stripped, re.IGNORECASE):
                return True
        return False
    
    def extract_question_number(self, text: str) -> Optional[int]:
        """Extract question number from text"""
        patterns = [
            r'Q(\d+)',
            r'Question\s+(\d+)',
            r'^(\d+)[\.\)]'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return int(match.group(1))
        return None
    
    def extract_marks(self, text: str) -> Optional[int]:
        """Extract marks from text (e.g., "[10 marks]", "(20 points)")"""
        patterns = [
            r'\[(\d+)\s*marks?\]',
            r'\((\d+)\s*marks?\)',
            r'\[(\d+)\s*points?\]',
            r'\((\d+)\s*points?\)',
            r'(\d+)\s*marks?',
            r'(\d+)\s*points?'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return int(match.group(1))
        return None
    
    def detect_question_type(self, question_text: str, options: List[str]) -> str:
        """Detect question type based on content"""
        # If no options, it's subjective
        if len(options) == 0:
            return 'subjective'
        
        question_lower = question_text.lower()
        
        # Check for keywords that indicate subjective questions
        if any(keyword in question_lower for keyword in ['explain', 'describe', 'discuss', 'compare', 'analyze', 'evaluate', 'what is', 'define']):
            return 'subjective'
        
        if any(keyword in question_lower for keyword in ['true', 'false', 't/f']):
            return 'true_false'
        
        # If has options, it's MCQ
        if len(options) > 0:
            return 'mcq'
        
        # Default to subjective for questions without options
        return 'subjective'
    
    def extract_questions(self) -> List[Dict[str, Any]]:
        """Extract questions from DOCX document"""
        questions = []
        current_question = None
        current_options = []
        in_question = False
        
        try:
            paragraphs = self.document.paragraphs
            
            for i, paragraph in enumerate(paragraphs):
                text = paragraph.text.strip()
                
                if not text:
                    continue
                
                # Check if this is a new question
                if self.is_question_start(text):
                    # Save previous question if exists
                    if current_question:
                        questions.append(current_question)
                    
                    # Start new question
                    question_num = self.extract_question_number(text) or len(questions) + 1
                    # Remove question prefix (Q1, Question 1, 1., etc.) and colon
                    question_text = re.sub(r'^(Q\d+|Question\s+\d+|\d+[\.\)])\s*:?\s*', '', text, flags=re.IGNORECASE).strip()
                    
                    # Extract marks if present
                    marks = self.extract_marks(text) or 10
                    
                    # Remove marks from question text
                    question_text = re.sub(r'\[(\d+)\s*marks?\]|\((\d+)\s*marks?\)', '', question_text, flags=re.IGNORECASE).strip()
                    
                    # Default to subjective if no options found later
                    current_question = {
                        'id': question_num,
                        'question': question_text,
                        'type': 'subjective',  # Default to subjective, will be updated if options found
                        'max_marks': marks,
                        'options': [],
                        'correct_answer': None,
                        'criteria': ['accuracy', 'completeness', 'clarity', 'logic', 'examples', 'structure']
                    }
                    
                    current_options = []
                    in_question = True
                    continue
                
                # If we're in a question, check for options or continuation
                if in_question and current_question:
                    # Check if this is an option (a), b), c), d), etc. or A., B., C., etc.
                    option_match = re.match(r'^([a-z]|[A-Z])[\.\)]\s*(.+)', text)
                    if option_match:
                        option_letter = option_match.group(1).upper()
                        option_text = option_match.group(2).strip()
                        
                        # Check if it's marked as correct (e.g., "a) Answer ✓" or "A) Answer (Correct)")
                        is_correct = '✓' in option_text or '(correct)' in option_text.lower() or '[correct]' in option_text.lower()
                        option_text = re.sub(r'[✓✔]|\(correct\)|\[correct\]', '', option_text, flags=re.IGNORECASE).strip()
                        
                        current_options.append({
                            'letter': option_letter,
                            'text': option_text,
                            'is_correct': is_correct
                        })
                        
                        # Change type to MCQ if options are found
                        if current_question['type'] == 'subjective':
                            current_question['type'] = 'mcq'
                            current_question['criteria'] = []
                        
                        if is_correct:
                            current_question['correct_answer'] = option_letter.lower()
                        
                        continue
                    
                    # Check if this continues the question text
                    if not current_options:
                        current_question['question'] += ' ' + text
                    else:
                        # Might be explanation or next question starting
                        if self.is_question_start(text):
                            # This is actually a new question
                            questions.append(current_question)
                            question_num = self.extract_question_number(text) or len(questions) + 1
                            # Remove question prefix (Q1, Question 1, 1., etc.) and colon
                            question_text = re.sub(r'^(Q\d+|Question\s+\d+|\d+[\.\)])\s*:?\s*', '', text, flags=re.IGNORECASE).strip()
                            marks = self.extract_marks(text) or 10
                            question_text = re.sub(r'\[(\d+)\s*marks?\]|\((\d+)\s*marks?\)', '', question_text, flags=re.IGNORECASE).strip()
                            
                            current_question = {
                                'id': question_num,
                                'question': question_text,
                                'type': 'subjective',  # Default to subjective
                                'max_marks': marks,
                                'options': [],
                                'correct_answer': None,
                                'criteria': ['accuracy', 'completeness', 'clarity', 'logic', 'examples', 'structure']
                            }
                            current_options = []
            
            # Save last question
            if current_question:
                # Set options
                current_question['options'] = current_options
                
                # Detect question type - if no options, it's subjective
                if len(current_options) == 0:
                    current_question['type'] = 'subjective'
                    current_question['criteria'] = ['accuracy', 'completeness', 'clarity', 'logic', 'examples', 'structure']
                else:
                    current_question['type'] = self.detect_question_type(
                        current_question['question'],
                        current_options
                    )
                
                questions.append(current_question)
            
            # Update question IDs to be sequential
            for i, q in enumerate(questions, 1):
                q['id'] = i
            
            logger.info(f"Extracted {len(questions)} questions")
            return questions
            
        except Exception as e:
            logger.error(f"Error extracting questions: {str(e)}")
            raise
    
    def validate_questions(self, questions: List[Dict[str, Any]]) -> bool:
        """Validate extracted questions"""
        if not questions:
            logger.error("No questions extracted from document")
            return False
        
        for i, q in enumerate(questions, 1):
            # Check question text
            question_text = q.get('question', '').strip()
            if not question_text:
                logger.error(f"Question {i} has no question text")
                return False
            
            question_type = q.get('type', 'subjective')
            options = q.get('options', [])
            
            # Auto-fix question type if needed
            if question_type == 'mcq' and len(options) == 0:
                logger.info(f"Question {i}: Auto-converting from MCQ to subjective (no options found)")
                q['type'] = 'subjective'
                q['criteria'] = ['accuracy', 'completeness', 'clarity', 'logic', 'examples', 'structure']
                q['correct_answer'] = None
                question_type = 'subjective'
            
            # Validate MCQ questions
            if question_type == 'mcq':
                if len(options) < 2:
                    logger.warning(f"Question {i} is MCQ but has insufficient options (need at least 2), converting to subjective")
                    q['type'] = 'subjective'
                    q['options'] = []
                    q['correct_answer'] = None
                    q['criteria'] = ['accuracy', 'completeness', 'clarity', 'logic', 'examples', 'structure']
                elif not q.get('correct_answer'):
                    logger.warning(f"Question {i} is MCQ but has no correct answer marked - will use first option as default")
                    # Use first option as default if no correct answer marked
                    if len(options) > 0:
                        q['correct_answer'] = options[0].get('letter', 'a').lower()
            
            # Ensure subjective questions have criteria
            if question_type == 'subjective' and not q.get('criteria'):
                q['criteria'] = ['accuracy', 'completeness', 'clarity', 'logic', 'examples', 'structure']
        
        logger.info(f"Validation passed for {len(questions)} questions")
        return True
    
    def convert_to_json(self) -> Dict[str, Any]:
        """Convert questions to JSON format"""
        try:
            metadata = self.extract_metadata()
            questions = self.extract_questions()
            
            if not self.validate_questions(questions):
                raise ValueError("Question validation failed")
            
            # Calculate total marks
            total_marks = sum(q.get('max_marks', 0) for q in questions)
            metadata['total_marks'] = total_marks if total_marks > 0 else metadata.get('total_marks', 0)
            
            quiz_json = {
                'metadata': metadata,
                'questions': questions,
                'total_questions': len(questions),
                'processed_at': datetime.now().isoformat()
            }
            
            logger.info("Successfully converted to JSON format")
            return quiz_json
            
        except Exception as e:
            logger.error(f"Error converting to JSON: {str(e)}")
            raise
    
    def save_json(self, quiz_json: Dict[str, Any]) -> str:
        """Save quiz JSON to file"""
        try:
            # Generate filename from source file
            source_name = os.path.splitext(os.path.basename(self.docx_path))[0]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{source_name}_{timestamp}.json"
            filepath = os.path.join(self.output_dir, filename)
            
            # Save JSON file
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(quiz_json, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Saved JSON to: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error saving JSON: {str(e)}")
            raise
    
    def process(self) -> tuple[int, str, str]:
        """
        Main processing function
        
        Returns:
            (exit_code, status_message, json_filepath)
        """
        try:
            # Validate file
            if not self.validate_file():
                return (1, "ERROR: File not found or invalid format", "")
            
            # Convert to JSON
            quiz_json = self.convert_to_json()
            
            # Save JSON
            json_filepath = self.save_json(quiz_json)
            
            # Success
            status = f"SUCCESS: Processed {len(quiz_json['questions'])} questions. Saved to {json_filepath}"
            logger.info(status)
            return (0, status, json_filepath)
            
        except ValueError as e:
            error_msg = f"ERROR: Validation failed - {str(e)}"
            logger.error(error_msg)
            return (2, error_msg, "")
            
        except Exception as e:
            error_msg = f"ERROR: Processing failed - {str(e)}"
            logger.error(error_msg)
            return (3, error_msg, "")


def main():
    """Main entry point"""
    if len(sys.argv) < 2:
        print("ERROR: Missing required argument: docx_file_path")
        print("Usage: python process_quiz_docx.py <docx_file_path> [output_dir]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "storage/quiz_json"
    
    # Process the file
    processor = QuizProcessor(docx_path, output_dir)
    exit_code, status, json_path = processor.process()
    
    # Print status (for PHP to capture)
    print(status)
    
    # Exit with appropriate code
    sys.exit(exit_code)


if __name__ == "__main__":
    main()

